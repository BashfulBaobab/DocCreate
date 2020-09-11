# Project DocMerge
# Author: Akshat Johari

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from os import listdir
from os.path import isfile, join
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from PIL import Image
import warnings

def main():
    print("Welcome to the document creation chamber.\nThis program creates .docx files with images and text, in tabular format.\n")
    name = input("Please enter the name you want this document to be given, without the .docx extension.\n")
    doc = create_doc(name)
    images = img_list()
    rows = input("Would you like the tables to have 6 images or 8?\n")
    while rows not in ("6", "8"):
        rows = input("Would you like the tables to have 6 images or 8?\nNote that if you enter less than 6 or 8, respectively, the table will create with the closest even number of rows that will fit the images.\n" )
    com = input("Do you want the ouput file to be compressed? Y/N\n")
    table_op(images, int(rows), doc, com)
    
    doc.save(name + ".docx")
    print("File " + name + " has been created in the current directory. Fare thee well.")
    
def create_doc(name):
    doc = Document()
    title = input("Please enter the title of your document.\n")
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Heading 1']
    p.style.font.size = Pt(16)
    p.style.font.bold = True
    
    sub = input("Does your document have a subheading?\nY for Yes, N for No\n")
    if sub in ("y", "Y"):
        sub = input("Please enter your subheading.\n")
        p = doc.add_paragraph(sub)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = doc.styles['Heading 2']
        p.style.font.size = Pt(14)
        p.style.paragraph_format.space_after = Pt(12)   
    
    return doc
              
def img_list():
    data_entry = input("How would you like to enter the images?\n1. Enter manually, one at a time\n2. Enter directory with all images (note that only jpg and png images are supported)\n3. Enter txt file with all images in separate lines\nPlease enter 1, 2, or 3: \t")
    while data_entry not in ("1", "2", "3"):
        print("Invalid entry. Please try again.")
        data_entry = input("How would you like to enter the images?\n1. Enter manually, one at a time\n2. Enter directory with all images (note that only jpg and png images are supported)\n3. Enter a .txt file with all images in separate lines\nPlease enter 1, 2, or 3: \t")
    #accept manual entries
    if data_entry == "1":
        l = []
        x = input("Enter image locations, either absolute location, or relative to current working directory. If you wish to stop data entry, type esc. Please make sure the locations are correct.\n")
        while x != "esc":
            l.append(x)
            x = input("Enter next image, or type escc to end input:\n")
        return l
    
    #accept directory with all images
    elif data_entry == "2":
        x = (".png", ".jpg")
        mypath = input("Enter directory:\n")
        l = [f for f in listdir(mypath) if (isfile(join(mypath, f)) and (f.endswith(".png") or f.endswith(".JPG")))]
        return l
    
    #accept txt file with list of images
    else:
        l = []
        f = input("Enter the .txt file:\n")
        f = open(f, 'r')
        while 1:
            line = f.readline()
            if not line:
                break
            l.append(line.strip())
        return l
    
def table_op(images, row_num, doc, com):
    
    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
    
    tbl = doc.add_table(rows = 1, cols = 2)
    tbl.style = 'TableGrid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #list is 0 indexed 
    counter = 0
    
    for i in range((len(images)+1)//2):
        if i%(row_num/2) == 0 and i!= 0:
            row = tbl.rows[0]
            remove_row(tbl, row)
            doc.add_page_break()
            tbl = doc.add_table(rows = 1, cols = 2)
            tbl.style = 'TableGrid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
        row_cells = tbl.add_row().cells
        #cell = row_cells[0]
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        i1 = images[counter]
        counter += 1
        if com == "y" or "Y":
            i1 = compressor(i1)
        if row_num/2 == 3:
            run.add_picture(i1, height = Inches(2))
        else:
            run.add_picture(i1, height = Inches(1.5))

        if counter != len(images):
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            i2 = images[counter]
            if com == "y" or "Y":
                i2 = compressor(i2)
            if row_num/2 == 3:
                run.add_picture(i2, height = Inches(2))
            else:
                run.add_picture(i2, height = Inches(1.5))
        tbl.add_row()
    
    row = tbl.rows[0]
    remove_row(tbl, row)
    
def compressor(img):
    x = Image.open(img)
    y = x.size
    r = tuple(int(round(i/4)) for i in y)
    x = x.resize(r, Image.ANTIALIAS)
    #saving image because docx and PIL Image are not compatible.
    x.save("scaled_img.jpg", optimize = True, quality = 95)
    return ("scaled_img.jpg")

if __name__ == "__main__":
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        main()